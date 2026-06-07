# ingestion/document_loader.py
# This file handles loading documents of different types
# It extracts text and metadata from PDF, DOCX, and TXT files

import os
from datetime import datetime
from pypdf import PdfReader
from docx import Document


def load_pdf(file_path):
    """
    Opens a PDF file and extracts all text from every page.
    Also collects metadata about the file.
    """
    print(f"Loading PDF: {file_path}")

    # Open the PDF file
    reader = PdfReader(file_path)

    # Extract text from every page and join them together
    full_text = ""
    for page_number, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            full_text += page_text + "\n"

    # Collect metadata about this document
    metadata = {
        "filename": os.path.basename(file_path),
        "file_type": "pdf",
        "file_path": file_path,
        "page_count": len(reader.pages),
        "uploaded_at": datetime.now().isoformat(),
        "file_size_bytes": os.path.getsize(file_path)
    }

    print(f"Successfully loaded PDF: {metadata['filename']} ({metadata['page_count']} pages)")
    return full_text, metadata


def load_docx(file_path):
    """
    Opens a Word document and extracts all text from every paragraph.
    Also collects metadata about the file.
    """
    print(f"Loading DOCX: {file_path}")

    # Open the Word document
    doc = Document(file_path)

    # Extract text from every paragraph and join them together
    full_text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text += paragraph.text + "\n"

    # Collect metadata about this document
    metadata = {
        "filename": os.path.basename(file_path),
        "file_type": "docx",
        "file_path": file_path,
        "paragraph_count": len(doc.paragraphs),
        "uploaded_at": datetime.now().isoformat(),
        "file_size_bytes": os.path.getsize(file_path)
    }

    print(f"Successfully loaded DOCX: {metadata['filename']} ({metadata['paragraph_count']} paragraphs)")
    return full_text, metadata


def load_txt(file_path):
    """
    Opens a plain text file and reads all content.
    Also collects metadata about the file.
    """
    print(f"Loading TXT: {file_path}")

    # Open and read the text file
    with open(file_path, "r", encoding="utf-8") as f:
        full_text = f.read()

    # Collect metadata about this document
    metadata = {
        "filename": os.path.basename(file_path),
        "file_type": "txt",
        "file_path": file_path,
        "uploaded_at": datetime.now().isoformat(),
        "file_size_bytes": os.path.getsize(file_path)
    }

    print(f"Successfully loaded TXT: {metadata['filename']}")
    return full_text, metadata


def load_document(file_path):
    """
    Master function that looks at the file extension
    and calls the right loader automatically.
    """
    # Get the file extension (like .pdf or .docx)
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return load_pdf(file_path)
    elif extension == ".docx":
        return load_docx(file_path)
    elif extension == ".txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


# This block runs only when you run this file directly
# It is used for testing
if __name__ == "__main__":
    print("Document Loader is ready!")
    print("Supported formats: PDF, DOCX, TXT")